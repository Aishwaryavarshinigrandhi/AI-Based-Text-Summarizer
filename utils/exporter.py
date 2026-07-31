import os
from datetime import datetime

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from docx import Document


OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "outputs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def export_summary_pdf(summary: str, filename: str) -> str:
    output_path = os.path.join(OUTPUT_DIR, f"{datetime.utcnow().strftime('%Y%m%d%H%M%S')}_{filename}.pdf")
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [Paragraph("AI Generated Summary", styles["Title"]), Spacer(1, 12), Paragraph(summary, styles["BodyText"])]
    doc.build(story)
    return output_path


def export_summary_docx(summary: str, filename: str) -> str:
    output_path = os.path.join(OUTPUT_DIR, f"{datetime.utcnow().strftime('%Y%m%d%H%M%S')}_{filename}.docx")
    document = Document()
    document.add_heading("AI Generated Summary", level=1)
    document.add_paragraph(summary)
    document.save(output_path)
    return output_path
